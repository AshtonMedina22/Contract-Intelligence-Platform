from __future__ import annotations

import io

from docx import Document

from lp_processor.models import NormalizedDocument, PdfPage
from lp_processor.parsers.base import DocumentParser

DOCX_MIME = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


class DocxParser(DocumentParser):
    """Native DOCX text extraction via python-docx. Not OCR."""

    parser_id = "docx-native"

    def supports(self, mime_type: str | None, filename: str | None) -> bool:
        name = (filename or "").lower()
        return name.endswith(".docx") or (mime_type or "") in DOCX_MIME

    def parse(self, payload: bytes, *, mime_type: str | None, filename: str | None) -> NormalizedDocument:
        document = Document(io.BytesIO(payload))
        warnings: list[str] = []
        pages: list[PdfPage] = []

        paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
        table_lines: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    table_lines.append(line)

        chunks: list[str] = []
        if paragraphs:
            chunks.append("\n".join(paragraphs))
        if table_lines:
            chunks.append("\n".join(table_lines))

        if not chunks:
            warnings.append("DOCX produced no paragraph or table text.")
            pages.append(PdfPage(page=1, text=""))
        else:
            # Approximate page breaks every ~3500 chars for source_page provenance.
            full = "\n\n".join(chunks)
            page_size = 3500
            for index, start in enumerate(range(0, max(len(full), 1), page_size), start=1):
                pages.append(PdfPage(page=index, text=full[start : start + page_size]))

        return NormalizedDocument(
            parser_id=self.parser_id,
            mime_type=mime_type,
            filename=filename,
            page_count=len(pages),
            sheet_count=None,
            sheets=[],
            pages=pages,
            warnings=warnings,
        )
